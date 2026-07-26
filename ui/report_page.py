# ui/report_page.py
import os
import io
import json
import time
import streamlit as st
from pathlib import Path
from datetime import datetime
from openai import OpenAI
from dotenv import load_dotenv

load_dotenv()

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# ── Load the sample report as the interpretation template ─────────────────────
TEMPLATE_PATH = Path(__file__).parent.parent / "knowledge_base" / "PLS-SEM_Results_Interpretation_Sample_Report.md"

def load_template() -> str:
    if TEMPLATE_PATH.exists():
        return TEMPLATE_PATH.read_text()
    # Fallback: inline summary of the format
    return """
    Format each section as: overview, reliability table + interpretation,
    structural paths table + interpretation, summary + recommendations.
    Be precise with numbers. Flag threshold violations clearly.
    """

# ─────────────────────────────────────────────────────────────────────────────
# OpenAI interpretation calls — one per section for focused, accurate output
# ─────────────────────────────────────────────────────────────────────────────

def call_openai(system_prompt: str, user_prompt: str) -> str:
    """Single OpenAI call with retry on rate limit."""
    for attempt in range(3):
        try:
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user",   "content": user_prompt}
                ],
                temperature=0.2,
                max_tokens=1200,
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            if attempt < 2:
                time.sleep(2 ** attempt)
            else:
                return f"[Generation failed: {e}]"


def generate_overview(stats: dict, template: str) -> str:
    constructs = [c["name"] for c in stats.get("constructs", [])]
    paths      = list(stats.get("bootstrapped_paths", {}).keys())
    meta       = stats.get("metadata", {})
    return call_openai(
        system_prompt=(
            "You are a PLS-SEM expert writing a formal academic report. "
            "Follow the structure and tone of the provided template exactly. "
            "Be concise, precise, and use hedged academic language where appropriate.\n\n"
            f"TEMPLATE REFERENCE:\n{template[:1500]}"
        ),
        user_prompt=(
            f"Write Section 1 (Overview) for a PLS-SEM report with these details:\n"
            f"- Sample size: {meta.get('n_obs', 'N/A')} observations\n"
            f"- Constructs: {', '.join(constructs)}\n"
            f"- Structural paths: {', '.join(paths)}\n"
            f"- Software: seminr v{meta.get('seminr_version', 'N/A')}\n\n"
            "Write 2-3 paragraphs explaining what this model tests and why "
            "measurement model assessment must precede structural interpretation. "
            "Do not include headers — plain paragraphs only."
        )
    )


def generate_reliability_interpretation(stats: dict, template: str) -> str:
    reliability = stats.get("reliability", {})
    rel_text    = json.dumps(reliability, indent=2)
    return call_openai(
        system_prompt=(
            "You are a PLS-SEM expert writing Section 2 of a formal report. "
            "Reference specific numbers from the data provided. "
            "Flag any construct that fails a threshold with a clear explanation of why it matters.\n\n"
            f"TEMPLATE REFERENCE (reliability section):\n{template[1500:3000]}"
        ),
        user_prompt=(
            f"Write the interpretation paragraph(s) for Section 2 "
            f"(Reliability & Convergent Validity) based on these results:\n\n"
            f"{rel_text}\n\n"
            "For each construct: state whether it passes or fails each metric "
            "(Cronbach alpha >= 0.70, CR >= 0.70, AVE >= 0.50, rhoA >= 0.70). "
            "Explain the practical implication of any failures. "
            "End with a recommendation about whether the measurement model is "
            "sound enough to proceed to structural interpretation. "
            "Do not include headers or tables — plain paragraphs only."
        )
    )


def generate_path_interpretation(stats: dict, template: str) -> str:
    boot_paths = stats.get("bootstrapped_paths", {})
    paths_text = json.dumps(boot_paths, indent=2)
    return call_openai(
        system_prompt=(
            "You are a PLS-SEM expert writing Section 3 of a formal report. "
            "Interpret bootstrapped path coefficients for a business audience. "
            "Always state the coefficient, t-statistic, p-value, and CI together.\n\n"
            f"TEMPLATE REFERENCE (structural paths section):\n{template[3000:4500]}"
        ),
        user_prompt=(
            f"Write the interpretation paragraph(s) for Section 3 "
            f"(Structural Path Results) based on these bootstrapped results:\n\n"
            f"{paths_text}\n\n"
            "For each path: state the coefficient direction and magnitude, "
            "whether t > 1.96, whether p < 0.05, whether the 95% CI excludes zero, "
            "and what this means in plain business language. "
            "If significant_ci is True, confirm significance via CI. "
            "Do not include headers or tables — plain paragraphs only."
        )
    )


def generate_r_squared_interpretation(stats: dict) -> str:
    r2       = stats.get("r_squared", {})
    f2       = stats.get("f_squared", {})
    q2       = stats.get("q_squared", {})
    combined = {"r_squared": r2, "f_squared": f2, "q_squared": q2}
    return call_openai(
        system_prompt=(
            "You are a PLS-SEM expert writing about model fit and predictive power. "
            "Use Hair et al. (2019) benchmarks: R2 >= 0.67 substantial, >= 0.33 moderate, "
            ">= 0.19 weak. f2: 0.02 small, 0.15 medium, 0.35 large. "
            "Q2 > 0 indicates predictive relevance."
        ),
        user_prompt=(
            f"Write 2-3 paragraphs interpreting these model fit statistics:\n\n"
            f"{json.dumps(combined, indent=2)}\n\n"
            "If q_squared is empty, note that blindfolding was not available "
            "in this seminr version and omit Q2 discussion. "
            "Translate each statistic into what it means for how well the model "
            "explains and predicts the outcome construct(s). "
            "Do not include headers — plain paragraphs only."
        )
    )


def generate_summary_recommendations(stats: dict, template: str) -> str:
    reliability   = stats.get("reliability", {})
    boot_paths    = stats.get("bootstrapped_paths", {})
    r2            = stats.get("r_squared", {})
    return call_openai(
        system_prompt=(
            "You are a PLS-SEM expert writing the final summary section of an academic report. "
            "Be direct about what passed, what failed, and what must happen next.\n\n"
            f"TEMPLATE REFERENCE (summary section):\n{template[4500:]}"
        ),
        user_prompt=(
            f"Write Section 4 (Summary & Recommendations) for a PLS-SEM report.\n\n"
            f"Reliability results: {json.dumps(reliability, indent=2)}\n"
            f"Path results: {json.dumps(boot_paths, indent=2)}\n"
            f"R-squared: {json.dumps(r2, indent=2)}\n\n"
            "Provide: (1) a 1-paragraph overall verdict on the measurement model, "
            "(2) a 1-paragraph verdict on the structural model, "
            "(3) a numbered list of 3-5 specific recommended next steps. "
            "Do not include headers — plain paragraphs followed by numbered list."
        )
    )


# ─────────────────────────────────────────────────────────────────────────────
# DOCX builder
# ─────────────────────────────────────────────────────────────────────────────

def build_docx(stats: dict, sections: dict) -> bytes:
    """Build the .docx report from stats tables + AI-generated interpretations."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    BLUE_DARK  = RGBColor(0x1E, 0x3A, 0x5F)
    BLUE_MID   = RGBColor(0x25, 0x63, 0xB0)
    GREEN      = RGBColor(0x16, 0xA3, 0x4A)
    RED        = RGBColor(0xDC, 0x26, 0x26)
    AMBER      = RGBColor(0xD9, 0x77, 0x06)
    MUTED      = RGBColor(0x64, 0x74, 0x8B)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)

    def set_cell_bg(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color)
        tcPr.append(shd)

    def add_heading(doc, text, level, color=None):
        p    = doc.add_heading(text, level=level)
        run  = p.runs[0] if p.runs else p.add_run(text)
        run.font.color.rgb = color or BLUE_DARK
        run.font.name = "Calibri"
        return p

    def add_body(doc, text, color=None, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.space_before = Pt(3)
        run = p.add_run(text)
        run.font.name    = "Calibri"
        run.font.size    = Pt(11)
        run.font.color.rgb = color or RGBColor(0x1A, 0x1A, 0x2E)
        run.font.italic  = italic
        return p

    def add_caption(doc, text):
        p   = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name    = "Calibri"
        run.font.size    = Pt(9)
        run.font.italic  = True
        run.font.color.rgb = MUTED
        p.paragraph_format.space_after = Pt(8)
        return p

    def header_row(table, headers, widths_cm):
        row = table.rows[0]
        for i, (hdr, w) in enumerate(zip(headers, widths_cm)):
            cell = row.cells[i]
            cell.width = Cm(w)
            set_cell_bg(cell, "1E3A5F")
            p   = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(hdr)
            run.font.bold      = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.name      = "Calibri"
            run.font.size      = Pt(10)

    def data_cell(cell, text, bold=False, color=None, bg=None, center=True):
        if bg:
            set_cell_bg(cell, bg)
        p   = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(str(text))
        run.font.name  = "Calibri"
        run.font.size  = Pt(10)
        run.font.bold  = bold
        if color:
            run.font.color.rgb = color

    def threshold_bg(value, threshold, higher_is_better=True):
        try:
            v  = float(value)
            ok = v >= threshold if higher_is_better else v < threshold
            return ("DCFCE7", GREEN) if ok else ("FEE2E2", RED)
        except:
            return ("FFFFFF", RGBColor(0,0,0))

    meta = stats.get("metadata", {})
    date = datetime.now().strftime("%d %B %Y")

    # ── Cover ──────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(80)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PLS-Assist")
    run.font.name  = "Calibri"
    run.font.size  = Pt(36)
    run.font.bold  = True
    run.font.color.rgb = BLUE_DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PLS-SEM Results Interpretation Report")
    run.font.name  = "Calibri"
    run.font.size  = Pt(18)
    run.font.color.rgb = BLUE_MID

    doc.add_paragraph()

    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_info = [
        ("Sample Size",    f"{meta.get('n_obs', 'N/A')} observations"),
        ("Constructs",     str(meta.get('n_constructs', 'N/A'))),
        ("Structural Paths", str(meta.get('n_paths', 'N/A'))),
        ("Software",       f"seminr v{meta.get('seminr_version', 'N/A')}"),
        ("Report Date",    date),
    ]
    for i, (label, value) in enumerate(meta_info):
        set_cell_bg(meta_table.rows[i].cells[0], "1E3A5F")
        r = meta_table.rows[i].cells[0].paragraphs[0].add_run(label)
        r.font.bold = True; r.font.color.rgb = RGBColor(255,255,255)
        r.font.name = "Calibri"; r.font.size = Pt(10)
        r2 = meta_table.rows[i].cells[1].paragraphs[0].add_run(value)
        r2.font.name = "Calibri"; r2.font.size = Pt(10)

    doc.add_page_break()

    # ── Section 1: Overview ────────────────────────────────────────────────────
    add_heading(doc, "1. Overview", 1)
    add_body(doc, sections["overview"])
    doc.add_paragraph()

    # ── Section 2: Reliability & Convergent Validity ──────────────────────────
    add_heading(doc, "2. Reliability & Convergent Validity", 1)

    reliability = stats.get("reliability", {})
    if reliability:
        add_caption(doc, "Table 2.1 — Reliability and Convergent Validity")
        tbl = doc.add_table(rows=len(reliability)+1, cols=6)
        tbl.style = "Table Grid"
        header_row(tbl, ["Construct", "Cronbach α", "CR (ρC)", "AVE", "ρA", "AVE Status"],
                   [4.0, 2.2, 2.2, 2.0, 2.0, 2.6])
        for i, (construct, vals) in enumerate(reliability.items(), 1):
            row   = tbl.rows[i]
            bg    = "F8FAFC" if i % 2 == 0 else "FFFFFF"
            alpha = vals.get("cronbach_alpha", 0)
            cr    = vals.get("composite_reliability", 0)
            ave   = vals.get("AVE", 0)
            rhoA  = vals.get("rho_A", 0)
            ave_bg, ave_color = threshold_bg(ave, 0.50)
            data_cell(row.cells[0], construct, center=False, bg=bg)
            data_cell(row.cells[1], f"{float(alpha):.3f}", color=GREEN if float(alpha)>=0.70 else RED, bg=bg)
            data_cell(row.cells[2], f"{float(cr):.3f}",    color=GREEN if float(cr)>=0.70 else RED,    bg=bg)
            data_cell(row.cells[3], f"{float(ave):.3f}",   bg=bg)
            data_cell(row.cells[4], f"{float(rhoA):.3f}",  bg=bg)
            data_cell(row.cells[5], "✓ >= 0.50" if float(ave)>=0.50 else "✗ < 0.50",
                      color=ave_color, bg=ave_bg, bold=True)

        add_caption(doc, "Thresholds: Cronbach α > 0.70; CR > 0.70; AVE > 0.50; ρA > 0.70. Source: Hair et al. (2019).")

    doc.add_paragraph()
    add_heading(doc, "Interpretation", 2)
    add_body(doc, sections["reliability"])
    doc.add_paragraph()

    # ── Section 3: Structural Paths ────────────────────────────────────────────
    add_heading(doc, "3. Structural Path Results (Bootstrapped)", 1)

    boot_paths = stats.get("bootstrapped_paths", {})
    if boot_paths:
        add_caption(doc, "Table 3.1 — Bootstrapped Path Coefficients")
        tbl2 = doc.add_table(rows=len(boot_paths)+1, cols=7)
        tbl2.style = "Table Grid"
        header_row(tbl2, ["Path", "β", "t-stat", "p-value", "2.5% CI", "97.5% CI", "Sig."],
                   [4.5, 1.5, 1.5, 1.8, 1.8, 1.8, 2.1])
        for i, (path_name, vals) in enumerate(boot_paths.items(), 1):
            row    = tbl2.rows[i]
            bg     = "F8FAFC" if i % 2 == 0 else "FFFFFF"
            orig   = vals.get("original")
            tstat  = vals.get("t_stat")
            pval   = vals.get("p_value")
            ci_lo  = vals.get("ci_lower")
            ci_hi  = vals.get("ci_upper")
            sig_ci = vals.get("significant_ci", False)
            sig    = (pval is not None and pval < 0.05) or sig_ci

            def fmt(v): return f"{float(v):.3f}" if v is not None else "—"
            def fmt_p(v):
                if v is None: return "—"
                return "< 0.001" if float(v) < 0.001 else f"{float(v):.3f}"

            data_cell(row.cells[0], path_name, center=False, bg=bg)
            data_cell(row.cells[1], fmt(orig),  bold=True, color=BLUE_MID, bg=bg)
            data_cell(row.cells[2], fmt(tstat), bg=bg)
            data_cell(row.cells[3], fmt_p(pval), color=GREEN if sig else RED, bg=bg)
            data_cell(row.cells[4], fmt(ci_lo),  bg=bg)
            data_cell(row.cells[5], fmt(ci_hi),  bg=bg)
            data_cell(row.cells[6], "✓ Sig." if sig else "✗ n.s.",
                      color=GREEN if sig else RED, bg="DCFCE7" if sig else "FEE2E2", bold=True)

        add_caption(doc, "β = standardised path coefficient. CI = 95% bias-corrected bootstrap CI (1,000 iterations).")

    doc.add_paragraph()
    add_heading(doc, "Interpretation", 2)
    add_body(doc, sections["paths"])
    doc.add_paragraph()

    # ── Section 4: R², f², Q² ─────────────────────────────────────────────────
    add_heading(doc, "4. Model Fit & Predictive Power", 1)

    r2 = stats.get("r_squared", {})
    if r2:
        add_caption(doc, "Table 4.1 — R² (Coefficient of Determination)")
        tbl3 = doc.add_table(rows=len(r2)+1, cols=3)
        tbl3.style = "Table Grid"
        header_row(tbl3, ["Construct", "R²", "Level"], [5.0, 2.5, 7.5])
        levels = {0.67: "Substantial", 0.33: "Moderate", 0.19: "Weak"}
        for i, (construct, val) in enumerate(r2.items(), 1):
            bg  = "F8FAFC" if i % 2 == 0 else "FFFFFF"
            v   = float(val) if val is not None else 0
            lvl = "Substantial" if v >= 0.67 else "Moderate" if v >= 0.33 else "Weak" if v >= 0.19 else "Very Weak"
            data_cell(tbl3.rows[i].cells[0], construct, center=False, bg=bg)
            data_cell(tbl3.rows[i].cells[1], f"{v:.3f}", bold=True, bg=bg)
            data_cell(tbl3.rows[i].cells[2], lvl,
                      color=GREEN if v>=0.33 else AMBER if v>=0.19 else RED, bg=bg)
        add_caption(doc, "Benchmarks: R² >= 0.67 substantial; >= 0.33 moderate; >= 0.19 weak (Hair et al., 2019).")

    doc.add_paragraph()
    add_heading(doc, "Interpretation", 2)
    add_body(doc, sections["r_squared"])
    doc.add_paragraph()

    # ── Section 5: Summary & Recommendations ──────────────────────────────────
    add_heading(doc, "5. Summary & Recommendations", 1)
    add_body(doc, sections["summary"])
    doc.add_paragraph()

    # ── References ─────────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "References", 1)
    refs = [
        "Hair, J. F., Risher, J. J., Sarstedt, M., & Ringle, C. M. (2019). When to use and how to report results of PLS-SEM. European Business Review, 31(1), 2–24.",
        "Henseler, J., Ringle, C. M., & Sarstedt, M. (2015). A new criterion for assessing discriminant validity in variance-based structural equation modeling. Journal of the Academy of Marketing Science, 43(1), 115–135.",
        "Dijkstra, T. K., & Henseler, J. (2015). Consistent and asymptotically normal PLS estimators for linear structural equations. Computational Statistics & Data Analysis, 81, 10–23.",
        f"Ray, S., & Danks, N. (2024). seminr: Building and Estimating Structural Equation Models. R package version {meta.get('seminr_version', '2.5.0')}.",
    ]
    for ref in refs:
        p   = doc.add_paragraph(style="List Bullet")
        run = p.add_run(ref)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(4)

    # ── Footer note ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    add_body(doc, f"Generated by PLS-Assist · {date} · Trinity College Dublin", color=MUTED, italic=True)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────────────────────────────────────
# Main page
# ─────────────────────────────────────────────────────────────────────────────

def show():
    from ui.upload_page import init_session_state
    init_session_state()

    st.title("📄 Interpretation Report")

    if not st.session_state["run_complete"]:
        st.warning(
            "No model has been estimated yet. "
            "Go to **Model Setup**, upload your dataset, define constructs and paths, "
            "then run the model."
        )
        return

    st.caption(
        "GPT-4o interprets your PLS-SEM results section-by-section using the "
        "PLS-SEM methodology knowledge base as a reference. "
        "Generation takes ~30–60 seconds."
    )

    # ── Generate button ───────────────────────────────────────────────────────
    if "report_sections" not in st.session_state:
        st.session_state["report_sections"] = None
    if "report_docx"     not in st.session_state:
        st.session_state["report_docx"]     = None

    col1, col2 = st.columns([2, 1])
    with col1:
        generate_btn = st.button(
            "Generate Report",
            type="primary",
            use_container_width=True,
            disabled=st.session_state["report_sections"] is not None
        )
    with col2:
        if st.button("Regenerate", use_container_width=True,
                     disabled=st.session_state["report_sections"] is None):
            st.session_state["report_sections"] = None
            st.session_state["report_docx"]     = None
            st.rerun()

    if generate_btn:
        stats    = st.session_state["model_stats"]
        template = load_template()

        progress = st.progress(0, text="Generating overview…")
        sections = {}

        sections["overview"]    = generate_overview(stats, template)
        progress.progress(20, text="Interpreting reliability & validity…")

        sections["reliability"] = generate_reliability_interpretation(stats, template)
        progress.progress(40, text="Interpreting structural paths…")

        sections["paths"]       = generate_path_interpretation(stats, template)
        progress.progress(60, text="Interpreting R², f², Q²…")

        sections["r_squared"]   = generate_r_squared_interpretation(stats)
        progress.progress(80, text="Generating summary & recommendations…")

        sections["summary"]     = generate_summary_recommendations(stats, template)
        progress.progress(90, text="Building Word document…")

        docx_bytes = build_docx(stats, sections)
        progress.progress(100, text="Done!")

        st.session_state["report_sections"] = sections
        st.session_state["report_docx"]     = docx_bytes
        st.rerun()

    # ── Display report ────────────────────────────────────────────────────────
    if st.session_state["report_sections"]:
        sections = st.session_state["report_sections"]
        stats    = st.session_state["model_stats"]

        # Download button at the top
        st.download_button(
            label="⬇️ Download Report (.docx)",
            data=st.session_state["report_docx"],
            file_name=f"PLS_Assist_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True
        )

        st.divider()

        # ── Section 1: Overview ───────────────────────────────────────────────
        st.header("1. Overview")
        st.write(sections["overview"])
        st.divider()

        # ── Section 2: Reliability ────────────────────────────────────────────
        st.header("2. Reliability & Convergent Validity")

        reliability = stats.get("reliability", {})
        if reliability:
            rel_rows = []
            for construct, vals in reliability.items():
                alpha = float(vals.get("cronbach_alpha", 0))
                cr    = float(vals.get("composite_reliability", 0))
                ave   = float(vals.get("AVE", 0))
                rhoA  = float(vals.get("rho_A", 0))
                rel_rows.append({
                    "Construct":  construct,
                    "Cronbach α": round(alpha, 3),
                    "CR (ρC)":    round(cr, 3),
                    "AVE":        round(ave, 3),
                    "ρA":         round(rhoA, 3),
                    "AVE Status": "✓ Acceptable" if ave >= 0.50 else "✗ Below 0.50",
                })
            import pandas as pd
            rel_df = pd.DataFrame(rel_rows).set_index("Construct")

            def color_ave(val):
                try:
                    return "color: green; font-weight: bold" if "✓" in str(val) else "color: red; font-weight: bold"
                except: return ""

            st.dataframe(rel_df.style.map(color_ave, subset=["AVE Status"]), use_container_width=True)
            st.caption("Thresholds: α > 0.70 · CR > 0.70 · AVE > 0.50 · ρA > 0.70")

        st.subheader("Interpretation")
        st.write(sections["reliability"])
        st.divider()

        # ── Section 3: Structural Paths ───────────────────────────────────────
        st.header("3. Structural Path Results (Bootstrapped)")

        boot_paths = stats.get("bootstrapped_paths", {})
        if boot_paths:
            import pandas as pd
            path_rows = []
            for path_name, vals in boot_paths.items():
                orig   = vals.get("original")
                tstat  = vals.get("t_stat")
                pval   = vals.get("p_value")
                ci_lo  = vals.get("ci_lower")
                ci_hi  = vals.get("ci_upper")
                sig_ci = vals.get("significant_ci", False)
                sig    = (pval is not None and float(pval) < 0.05) or sig_ci

                def fmt(v): return round(float(v), 3) if v is not None else "—"
                def fmt_p(v):
                    if v is None: return "—"
                    return "< 0.001" if float(v) < 0.001 else round(float(v), 3)

                path_rows.append({
                    "Path":       path_name,
                    "β":          fmt(orig),
                    "t-stat":     fmt(tstat),
                    "p-value":    fmt_p(pval),
                    "2.5% CI":    fmt(ci_lo),
                    "97.5% CI":   fmt(ci_hi),
                    "Significant": "✓ Yes" if sig else "✗ No",
                })
            path_df = pd.DataFrame(path_rows).set_index("Path")

            def color_sig(val):
                return "color: green; font-weight: bold" if "✓" in str(val) else "color: red; font-weight: bold"

            st.dataframe(path_df.style.map(color_sig, subset=["Significant"]), use_container_width=True)
            st.caption("β = standardised path coefficient. 95% bias-corrected bootstrap CI (1,000 iterations).")

        st.subheader("Interpretation")
        st.write(sections["paths"])
        st.divider()

        # ── Section 4: R², f², Q² ─────────────────────────────────────────────
        st.header("4. Model Fit & Predictive Power")
        r2 = stats.get("r_squared", {})
        if r2:
            import pandas as pd
            r2_rows = []
            for construct, val in r2.items():
                v   = float(val) if val is not None else 0
                lvl = "Substantial" if v >= 0.67 else "Moderate" if v >= 0.33 else "Weak" if v >= 0.19 else "Very Weak"
                r2_rows.append({"Construct": construct, "R²": round(v, 3), "Level": lvl})
            st.dataframe(pd.DataFrame(r2_rows).set_index("Construct"), use_container_width=True)
            st.caption("Benchmarks: R² >= 0.67 substantial · >= 0.33 moderate · >= 0.19 weak (Hair et al., 2019).")

        st.subheader("Interpretation")
        st.write(sections["r_squared"])
        st.divider()

        # ── Section 5: Summary ────────────────────────────────────────────────
        st.header("5. Summary & Recommendations")
        st.write(sections["summary"])
        st.divider()

        # Download button again at the bottom
        st.download_button(
            label="⬇️ Download Report (.docx)",
            data=st.session_state["report_docx"],
            file_name=f"PLS_Assist_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )